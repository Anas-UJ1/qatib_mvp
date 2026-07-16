import os
import re
import fitz  # PyMuPDF
from typing import List
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
import logging

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Security: hard cap on uploaded PDF size to avoid resource exhaustion
# during a live demo (malicious or accidentally huge file).
MAX_UPLOAD_SIZE_MB = int(os.environ.get("MAX_UPLOAD_SIZE_MB", 25))
MAX_UPLOAD_SIZE_BYTES = MAX_UPLOAD_SIZE_MB * 1024 * 1024


class RegulatoryDocumentParser:
    def __init__(self, chunk_size: int = 800, chunk_overlap: int = 150):
        # Arabic legal text is dense and structured around numbered articles
        # ("المادة 12") and clauses ("البند"). We split on those boundaries
        # FIRST, before falling back to generic separators, so a chunk is
        # far less likely to straddle two unrelated articles.
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=[
                r"(?=المادة\s*\d+)",   # split before "المادة <N>" (Article N)
                r"(?=البند\s*\d+)",    # split before "البند <N>" (Clause N)
                r"(?=الفقرة\s*\d+)",   # split before "الفقرة <N>" (Paragraph N)
                "\n\n",
                "\n",
                ".",
                "؟",
                " ",
                "",
            ],
            is_separator_regex=True,
        )

    def clean_text(self, text: str) -> str:
        if not text:
            return ""
        text = re.sub(r"\n+", "\n", text)
        text = re.sub(r"\s{2,}", " ", text)
        return text.strip()

    def _check_size_bytes(self, size_bytes: int, display_name: str) -> bool:
        """Returns True if the given byte count is within the allowed limit."""
        if size_bytes > MAX_UPLOAD_SIZE_BYTES:
            size_mb = size_bytes / (1024 * 1024)
            logger.warning(
                f"Rejected '{display_name}': "
                f"{size_mb:.1f}MB exceeds the {MAX_UPLOAD_SIZE_MB}MB limit."
            )
            return False
        return True

    def _check_file_size(self, file_path: str) -> bool:
        """Returns True if file is within the allowed size limit."""
        try:
            size_bytes = os.path.getsize(file_path)
        except OSError as e:
            logger.error(f"Could not stat file {file_path}: {str(e)}")
            return False
        return self._check_size_bytes(size_bytes, os.path.basename(file_path))

    def extract_text(self, file_path: str) -> List[Document]:
        documents = []

        # --- Security: reject oversized files before opening with fitz ---
        if not self._check_file_size(file_path):
            return []

        try:
            doc = fitz.open(file_path)
            file_name = os.path.basename(file_path)
            logger.info(f"Parsing {file_name} - Total Pages: {len(doc)}")
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text = page.get_text("text")
                cleaned_text = self.clean_text(text)
                if cleaned_text:
                    metadata = {
                        "source": file_name,
                        "page_number": page_num + 1,
                        "regulatory_body": self._infer_regulatory_body(file_name),
                    }
                    documents.append(Document(page_content=cleaned_text, metadata=metadata))
            doc.close()
            return documents
        except Exception as e:
            logger.error(f"Error parsing {file_path}: {str(e)}")
            return []

    def extract_text_from_pdf_bytes(self, file_bytes: bytes, display_name: str) -> List[Document]:
        """Parse an in-memory PDF upload (e.g. a Streamlit UploadedFile's
        bytes). Unlike extract_text(), this never touches disk and tags
        chunks as doc_kind='contract' rather than inferring a
        regulatory_body -- an uploaded contract is the subject being
        reviewed, not a regulation source, and must never be indexed into
        the shared qatib_regulations collection."""
        if not self._check_size_bytes(len(file_bytes), display_name):
            return []

        # --- Security: validate actual content, not just the .pdf extension ---
        if file_bytes[:5] != b"%PDF-":
            logger.error(f"Rejected '{display_name}': not a valid PDF (magic bytes mismatch).")
            return []

        try:
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            documents = []
            logger.info(f"Parsing uploaded PDF '{display_name}' - Total Pages: {len(doc)}")
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                cleaned_text = self.clean_text(page.get_text("text"))
                if cleaned_text:
                    documents.append(Document(
                        page_content=cleaned_text,
                        metadata={
                            "source": display_name,
                            "page_number": page_num + 1,
                            "doc_kind": "contract",
                        },
                    ))
            doc.close()
            return documents
        except Exception as e:
            logger.error(f"Error parsing uploaded PDF '{display_name}': {str(e)}")
            return []

    def extract_text_from_docx_bytes(self, file_bytes: bytes, display_name: str) -> List[Document]:
        """Parse an in-memory DOCX upload. DOCX (OOXML) files are ZIP
        containers, so a 'PK' magic-byte check plus a real parse attempt
        (which fails on a malformed/renamed non-docx zip) stands in for
        proper content validation."""
        if not self._check_size_bytes(len(file_bytes), display_name):
            return []

        if file_bytes[:2] != b"PK":
            logger.error(f"Rejected '{display_name}': not a valid DOCX (magic bytes mismatch).")
            return []

        try:
            import io
            from docx import Document as DocxDocument
            docx_obj = DocxDocument(io.BytesIO(file_bytes))
        except Exception as e:
            logger.error(f"Error opening uploaded DOCX '{display_name}': {str(e)}")
            return []

        # DOCX stores no pagination -- page breaks are a print-layout
        # concern Word computes at render time, not persisted in the XML.
        # Synthesize sequential ~1800-char "sections" as a page_number
        # surrogate so citations still have something concrete to point to.
        SECTION_CHAR_SIZE = 1800
        documents, buffer, section_no = [], [], 1
        for para in docx_obj.paragraphs:
            text = self.clean_text(para.text)
            if not text:
                continue
            buffer.append(text)
            if sum(len(t) for t in buffer) >= SECTION_CHAR_SIZE:
                documents.append(Document(
                    page_content="\n".join(buffer),
                    metadata={"source": display_name, "page_number": section_no, "doc_kind": "contract"},
                ))
                buffer, section_no = [], section_no + 1
        if buffer:
            documents.append(Document(
                page_content="\n".join(buffer),
                metadata={"source": display_name, "page_number": section_no, "doc_kind": "contract"},
            ))

        logger.info(f"Parsed uploaded DOCX '{display_name}' - {len(documents)} section(s).")
        return documents

    def _infer_regulatory_body(self, file_name: str) -> str:
        name_lower = file_name.lower()
        if "sama" in name_lower:
            return "SAMA"
        if "zatca" in name_lower or "vat" in name_lower or "zakat" in name_lower:
            return "ZATCA"
        if "cma" in name_lower or "capital" in name_lower:
            return "CMA"
        if "sdaia" in name_lower or "personal data" in name_lower:
            return "SDAIA"
        return "General Regulation"

    def process_document(self, file_path: str) -> List[Document]:
        logger.info(f"Starting processing pipeline for: {file_path}")
        raw_docs = self.extract_text(file_path)
        if not raw_docs:
            return []
        chunked_docs = self.text_splitter.split_documents(raw_docs)
        logger.info(
            f"Successfully created {len(chunked_docs)} chunks from {os.path.basename(file_path)}"
        )
        return chunked_docs
